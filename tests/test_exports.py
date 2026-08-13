from __future__ import annotations

import tempfile
import builtins
import importlib.util
import sys
import types
import unittest
import zipfile
from decimal import Decimal
from pathlib import Path
from unittest.mock import Mock, patch

from docx import Document
from openpyxl import load_workbook


def _load_exports_module():
    root = Path(__file__).resolve().parents[1]
    package_name = "_exports_test_pkg"
    package = types.ModuleType(package_name)
    package.__path__ = [str(root / "text2sql_agent")]
    sys.modules[package_name] = package

    config = types.ModuleType(f"{package_name}.config")
    config.REPORT_DIR = root / "reports"
    config.MAX_QUERY_ROW_LIMIT = 1_000_000
    config.EXPORT_QUERY_TIMEOUT_MS = 300_000
    sys.modules[f"{package_name}.config"] = config

    spec = importlib.util.spec_from_file_location(
        f"{package_name}.exports",
        root / "text2sql_agent" / "exports.py",
    )
    if spec is None or spec.loader is None:
        raise RuntimeError("exports.py를 불러올 수 없습니다.")
    module = importlib.util.module_from_spec(spec)
    sys.modules[f"{package_name}.exports"] = module
    spec.loader.exec_module(module)
    return module


exports = _load_exports_module()


class ExportReportTest(unittest.TestCase):
    def setUp(self) -> None:
        self.result = {
            "question": "2026년 월별 법인 매출 알려줘",
            "answer": "2026년 1월 매출은 1,234만원이고 2월 매출은 2,345만원입니다.",
            "query_columns": ["월", "매출금액", "거래건수"],
            "query_rows": [
                ("2026-01", Decimal("12340000"), 120),
                ("2026-02", Decimal("23450000"), 210),
            ],
            "final_sql": "SELECT month, amount, count FROM sample_sales",
            "matched_query_name": "월별 매출",
        }

    def test_export_all_creates_openable_word_excel_and_text_reports(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            with patch.object(exports, "REPORT_DIR", Path(tmpdir)):
                paths = exports.export_all(self.result)

                self.assertEqual({"word", "excel", "text"}, set(paths))
                for path in paths.values():
                    self.assertTrue(Path(path).exists(), path)

                document = Document(paths["word"])
                word_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                self.assertIn("KB카드 법인영업 데이터 분석 보고서", word_text)
                self.assertIn("분석 결과", word_text)

                workbook = load_workbook(paths["excel"], data_only=True)
                self.assertEqual(["요약", "상세 데이터", "SQL"], workbook.sheetnames)
                summary = workbook["요약"]
                self.assertEqual("KB카드 법인영업 데이터 분석 보고서", summary["A1"].value)
                self.assertEqual("질문", summary["A4"].value)
                self.assertEqual(self.result["question"], summary["B4"].value)

                data_sheet = workbook["상세 데이터"]
                self.assertEqual(["월", "매출금액", "거래건수"], [cell.value for cell in data_sheet[1]])
                self.assertEqual("2026-01", data_sheet["A2"].value)
                self.assertEqual(12340000, data_sheet["B2"].value)
                self.assertEqual(120, data_sheet["C2"].value)

                text = Path(paths["text"]).read_text(encoding="utf-8")
                self.assertIn("[분석 결과]", text)
                self.assertIn("[상세 데이터] (총 2건)", text)
                self.assertIn("[실행 SQL]", text)

    def test_excel_export_handles_dict_rows_without_query_columns(self) -> None:
        result = {
            "question": "딕셔너리 결과 저장",
            "answer": "딕셔너리 행도 저장됩니다.",
            "query_rows": [
                {"월": "2026-01", "매출금액": 1000},
                {"월": "2026-02", "매출금액": 2000},
            ],
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            with patch.object(exports, "REPORT_DIR", Path(tmpdir)):
                path = exports.export_to_excel(result)

                workbook = load_workbook(path, data_only=True)
                data_sheet = workbook["상세 데이터"]
                self.assertEqual(["월", "매출금액"], [cell.value for cell in data_sheet[1]])
                self.assertEqual("2026-02", data_sheet["A3"].value)
                self.assertEqual(2000, data_sheet["B3"].value)

    def test_prepare_export_result_reloads_all_rows_without_mutating_preview(self) -> None:
        db_module = types.ModuleType(f"{exports.__package__}.db")
        execute_sql = Mock(
            return_value=(["값"], [(1,), (2,), (3,)], None)
        )
        db_module.execute_sql = execute_sql
        preview = {
            "final_sql": "SELECT value FROM sample LIMIT 1000000",
            "query_columns": ["값"],
            "query_rows": [(1,)],
        }

        with patch.dict(sys.modules, {f"{exports.__package__}.db": db_module}):
            prepared = exports.prepare_export_result(preview)

        self.assertEqual([(1,)], preview["query_rows"])
        self.assertEqual([(1,), (2,), (3,)], prepared["query_rows"])
        execute_sql.assert_called_once_with(
            preview["final_sql"],
            max_rows=1_000_000,
            statement_timeout_ms=300_000,
        )

    def test_prepare_export_result_preserves_current_snapshot_fallback_decision(self) -> None:
        db_module = types.ModuleType(f"{exports.__package__}.db")
        execute_sql = Mock(return_value=(["값"], [(1,)], None))
        db_module.execute_sql = execute_sql
        workflow_module = types.ModuleType(f"{exports.__package__}.workflow")
        allow_fallback = Mock(return_value=False)
        workflow_module._allow_cross_cycle_fallback = allow_fallback
        preview = {
            "question": "현재 가맹점 현황",
            "final_sql": (
                'SELECT a."기준년월일" FROM card_system.tbdaaus01 a '
                'WHERE a."기준년월일" = \'20260811\''
            ),
            "query_columns": ["값"],
            "query_rows": [(1,)],
        }

        with patch.dict(
            sys.modules,
            {
                f"{exports.__package__}.db": db_module,
                f"{exports.__package__}.workflow": workflow_module,
            },
        ):
            exports.prepare_export_result(preview)

        allow_fallback.assert_called_once_with(preview["question"], preview["final_sql"])
        execute_sql.assert_called_once_with(
            preview["final_sql"],
            max_rows=1_000_000,
            statement_timeout_ms=300_000,
            allow_cross_cycle_fallback=False,
        )

    def test_large_excel_export_uses_write_only_path_and_keeps_every_row(self) -> None:
        result = {
            "question": "대용량 조회",
            "query_columns": ["순번", "값"],
            "query_rows": [(1, "가"), (2, "나"), (3, "다")],
            "final_sql": "SELECT sequence, value FROM sample LIMIT 1000000",
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            with (
                patch.object(exports, "REPORT_DIR", Path(tmpdir)),
                patch.object(exports, "_LARGE_EXCEL_ROW_THRESHOLD", 2),
            ):
                path = exports.export_to_excel(result)

            workbook = load_workbook(path, data_only=True)
            self.assertEqual(["요약", "상세 데이터", "SQL"], workbook.sheetnames)
            data_rows = list(workbook["상세 데이터"].iter_rows(values_only=True))
            self.assertEqual(("순번", "값"), data_rows[0])
            self.assertEqual([(1, "가"), (2, "나"), (3, "다")], data_rows[1:])
            self.assertEqual("A1:B4", workbook["상세 데이터"].auto_filter.ref)

    def test_text_export_keeps_rows_beyond_the_report_preview_size(self) -> None:
        result = {
            "question": "TXT 전체 데이터",
            "query_columns": ["순번"],
            "query_rows": [(index,) for index in range(250)],
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            with patch.object(exports, "REPORT_DIR", Path(tmpdir)):
                path = exports.export_to_text(result)

            text = Path(path).read_text(encoding="utf-8")
            self.assertIn("[상세 데이터] (총 250건)", text)
            self.assertIn("\n249\n", text)
            self.assertNotIn("상위 200건만 표시", text)

    def test_word_export_falls_back_when_python_docx_is_missing(self) -> None:
        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == "docx" or name.startswith("docx."):
                raise ImportError("docx unavailable")
            return real_import(name, *args, **kwargs)

        with tempfile.TemporaryDirectory() as tmpdir:
            with patch.object(exports, "REPORT_DIR", Path(tmpdir)):
                with patch("builtins.__import__", side_effect=fake_import):
                    path = exports.export_to_word(self.result)

                self.assertTrue(Path(path).exists(), path)
                with zipfile.ZipFile(path) as docx:
                    self.assertIn("word/document.xml", docx.namelist())
                    document_xml = docx.read("word/document.xml").decode("utf-8")
                self.assertIn("KB카드 법인영업 데이터 분석 보고서", document_xml)
                self.assertIn("상세 데이터", document_xml)

                document = Document(path)
                word_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                self.assertIn("KB카드 법인영업 데이터 분석 보고서", word_text)


if __name__ == "__main__":
    unittest.main()
