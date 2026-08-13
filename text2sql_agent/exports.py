"""Report/export helpers for completed query results."""

import csv
import re
import zipfile
from datetime import date, datetime
from decimal import Decimal
from html import escape
from typing import Any

from .config import REPORT_DIR

# ---------------------------------------------------------------------------
# 7. 보고서 내보내기
# ---------------------------------------------------------------------------


def _ensure_report_dir():
    REPORT_DIR.mkdir(exist_ok=True)


def _make_filename(question: str, ext: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = re.sub(r"[^\w가-힣]", "_", question)[:30].strip("_")
    return f"{ts}_{safe}.{ext}"


def format_number_for_report(val) -> str:
    if isinstance(val, (int, float, Decimal)):
        val = float(val)
        if abs(val) >= 100_000_000:
            return f"{val / 100_000_000:,.1f}억"
        elif abs(val) >= 10_000:
            return f"{val / 10_000:,.0f}만"
        else:
            return f"{val:,}"
    return str(val)


def _is_numeric_value(value: Any) -> bool:
    return isinstance(value, (int, float, Decimal)) and not isinstance(value, bool)


def _row_values(row: Any, columns: list[Any] | None = None) -> list[Any]:
    if isinstance(row, dict):
        if columns:
            return [row.get(col, row.get(str(col), "")) for col in columns]
        return list(row.values())
    if isinstance(row, (list, tuple)):
        values = list(row)
    else:
        values = [row]
    if columns:
        if len(values) < len(columns):
            values.extend([""] * (len(columns) - len(values)))
        return values[: len(columns)]
    return values


def _raw_result_columns_and_rows(result: dict) -> tuple[list[Any], list[Any]]:
    raw_rows = result.get("query_rows", []) or []
    raw_columns = list(result.get("query_columns", []) or [])
    if not raw_columns and raw_rows:
        first_row = raw_rows[0]
        if isinstance(first_row, dict):
            raw_columns = list(first_row.keys())
        else:
            raw_columns = [f"컬럼{i + 1}" for i in range(len(_row_values(first_row)))]
    return raw_columns, raw_rows


def _result_columns_and_rows(result: dict, row_limit: int | None = None) -> tuple[list[str], list[list[Any]]]:
    raw_columns, raw_rows = _raw_result_columns_and_rows(result)
    selected_rows = raw_rows if row_limit is None else raw_rows[:row_limit]
    columns = [str(col) for col in raw_columns]
    rows = [_row_values(row, raw_columns) for row in selected_rows]
    return columns, rows


def prepare_export_result(result: dict) -> dict:
    """Reload the complete SQL result for file export without growing UI/session payloads."""

    prepared = dict(result)
    sql = str(result.get("final_sql") or "").strip()
    if not sql or result.get("bad_debt_excel_path") or result.get("selected_tool") == "대손비용률_분석":
        return prepared

    from .config import EXPORT_QUERY_TIMEOUT_MS, MAX_QUERY_ROW_LIMIT
    from .db import execute_sql

    execute_kwargs = {
        "max_rows": MAX_QUERY_ROW_LIMIT,
        "statement_timeout_ms": EXPORT_QUERY_TIMEOUT_MS,
    }
    question = str(result.get("question") or "").strip()
    if question:
        # 화면 조회에서 최신 D-1 스냅샷의 월별 폴백을 막았다면
        # 다운로드용 전체 재조회도 같은 의도를 유지해야 한다. 순환 import를
        # 피하기 위해 실행 시점에만 workflow의 단일 판정 함수를 재사용한다.
        from .workflow import _allow_cross_cycle_fallback

        if not _allow_cross_cycle_fallback(question, sql):
            execute_kwargs["allow_cross_cycle_fallback"] = False

    columns, rows, error = execute_sql(sql, **execute_kwargs)
    if error:
        raise RuntimeError("다운로드용 전체 데이터를 조회하지 못했습니다.") from None

    if str(result.get("followup_mode") or "") in {"transform", "transform_visualization"}:
        from .followup_ops import apply_local_transform

        transformed = apply_local_transform(str(result.get("followup_question") or ""), columns, rows)
        columns, rows = transformed["columns"], transformed["rows"]

    prepared["query_columns"] = columns
    prepared["query_rows"] = rows
    return prepared


def _excel_cell_value(value: Any) -> Any:
    if isinstance(value, Decimal):
        return float(value)
    if isinstance(value, (datetime, date, str, int, float, bool)) or value is None:
        return value
    return str(value)


def _word_run(value: Any, *, bold: bool = False, size: int = 20, color: str | None = None) -> str:
    props = [
        '<w:rFonts w:ascii="Malgun Gothic" w:hAnsi="Malgun Gothic" w:eastAsia="Malgun Gothic"/>',
        f'<w:sz w:val="{size}"/>',
    ]
    if bold:
        props.append("<w:b/>")
    if color:
        props.append(f'<w:color w:val="{color}"/>')
    text = str(value or "")
    parts = text.splitlines() or [""]
    runs = []
    for i, part in enumerate(parts):
        if i:
            runs.append("<w:r><w:br/></w:r>")
        runs.append(
            "<w:r><w:rPr>"
            + "".join(props)
            + '</w:rPr><w:t xml:space="preserve">'
            + escape(part, quote=False)
            + "</w:t></w:r>"
        )
    return "".join(runs)


def _word_paragraph(value: Any, *, bold: bool = False, size: int = 20, color: str | None = None, align: str | None = None) -> str:
    align_xml = f'<w:jc w:val="{align}"/>' if align else ""
    return f"<w:p><w:pPr>{align_xml}</w:pPr>{_word_run(value, bold=bold, size=size, color=color)}</w:p>"


def _word_cell(value: Any, *, bold: bool = False, shaded: bool = False, align: str | None = None) -> str:
    shade_xml = '<w:shd w:fill="D9EAF7"/>' if shaded else ""
    align_xml = f'<w:vAlign w:val="center"/>' if align == "center" else ""
    paragraph_align = align if align in {"center", "right"} else None
    return (
        "<w:tc><w:tcPr>"
        '<w:tcW w:w="2400" w:type="dxa"/>'
        + shade_xml
        + align_xml
        + "</w:tcPr>"
        + _word_paragraph(value, bold=bold, size=18, align=paragraph_align)
        + "</w:tc>"
    )


def _word_table(headers: list[str], rows: list[list[Any]]) -> str:
    border = (
        "<w:tblBorders>"
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D9E2EC"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="D9E2EC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D9E2EC"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="D9E2EC"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D9E2EC"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D9E2EC"/>'
        "</w:tblBorders>"
    )
    table_rows = ["<w:tr>" + "".join(_word_cell(header, bold=True, shaded=True, align="center") for header in headers) + "</w:tr>"]
    for row in rows:
        cells = []
        for value in row:
            align = "right" if _is_numeric_value(value) else None
            cells.append(_word_cell(format_number_for_report(value), align=align))
        table_rows.append("<w:tr>" + "".join(cells) + "</w:tr>")
    return "<w:tbl><w:tblPr>" + border + "</w:tblPr>" + "".join(table_rows) + "</w:tbl>"


def _export_to_word_fallback(result: dict, filepath) -> str:
    question = result.get("question", "조회결과")
    total_rows = len(result.get("query_rows", []) or [])
    columns, rows = _result_columns_and_rows(result, row_limit=200)
    answer = result.get("answer", "")
    sql = result.get("final_sql", "")

    body = [
        _word_paragraph("KB카드 법인영업 데이터 분석 보고서", bold=True, size=32, align="center"),
        _word_paragraph(""),
        _word_table(
            ["항목", "내용"],
            [
                ["작성일시", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                ["질문", question],
                ["분석 경로", _get_source_label(result)],
            ],
        ),
    ]
    if answer:
        body.extend([
            _word_paragraph("분석 결과", bold=True, size=26, color="2F5597"),
            _word_paragraph(answer),
        ])
    if columns and rows:
        body.extend([
            _word_paragraph("상세 데이터", bold=True, size=26, color="2F5597"),
            _word_paragraph(f"총 {total_rows}건 조회됨"),
            _word_table(columns, rows),
        ])
        if total_rows > 200:
            body.append(_word_paragraph(f"* 전체 {total_rows}건 중 상위 200건만 표시", size=18))
    if sql:
        body.extend([
            _word_paragraph("실행 SQL", bold=True, size=26, color="2F5597"),
            _word_paragraph(sql, size=16, color="505050"),
        ])
    body.extend([
        _word_paragraph(""),
        _word_paragraph("KB카드 법인영업 Text2SQL Agent v10 자동생성 보고서", size=16, color="969696", align="center"),
    ])

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + "".join(body)
        + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/>'
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>'
        "</w:sectPr></w:body></w:document>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )

    with zipfile.ZipFile(filepath, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types)
        docx.writestr("_rels/.rels", rels)
        docx.writestr("word/document.xml", document_xml)
    return str(filepath)


def export_to_word(result: dict) -> str:
    _ensure_report_dir()
    question = result.get("question", "조회결과")
    filepath = REPORT_DIR / _make_filename(question, "docx")
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return _export_to_word_fallback(result, filepath)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    doc.add_heading("KB카드 법인영업 데이터 분석 보고서", level=0)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Light List"
    for i, (label, value) in enumerate([
        ("작성일시", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("질문", question),
        ("분석 경로", _get_source_label(result)),
    ]):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    doc.add_paragraph("")
    answer = result.get("answer", "")
    if answer:
        doc.add_heading("분석 결과", level=1)
        doc.add_paragraph(answer)
    total_rows = len(result.get("query_rows", []) or [])
    columns, rows = _result_columns_and_rows(result, row_limit=200)
    if columns and rows:
        doc.add_heading("상세 데이터", level=1)
        doc.add_paragraph(f"총 {total_rows}건 조회됨")
        table = doc.add_table(rows=1 + len(rows), cols=len(columns))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, col_name in enumerate(columns):
            cell = table.rows[0].cells[j]
            cell.text = col_name
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = format_number_for_report(val)
                for paragraph in cell.paragraphs:
                    if _is_numeric_value(val):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(9)
        if total_rows > 200:
            doc.add_paragraph(f"* 전체 {total_rows}건 중 상위 200건만 표시")
    sql = result.get("final_sql", "")
    if sql:
        doc.add_heading("실행 SQL", level=1)
        sql_para = doc.add_paragraph()
        sql_run = sql_para.add_run(sql)
        sql_run.font.size = Pt(8)
        sql_run.font.name = "Consolas"
        sql_run.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph("")
    footer = doc.add_paragraph("KB카드 법인영업 Text2SQL Agent v10 자동생성 보고서")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)
    doc.save(str(filepath))
    return str(filepath)


_LARGE_EXCEL_ROW_THRESHOLD = 100_000
_EXCEL_MAX_DATA_ROWS = 1_048_575


def _export_to_excel_write_only(result: dict, filepath) -> str:
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter

    raw_columns, raw_rows = _raw_result_columns_and_rows(result)
    if len(raw_rows) > _EXCEL_MAX_DATA_ROWS:
        raise ValueError(f"Excel은 헤더를 제외하고 최대 {_EXCEL_MAX_DATA_ROWS:,}행까지 저장할 수 있습니다.")

    columns = [str(column) for column in raw_columns]
    wb = Workbook(write_only=True)
    summary = wb.create_sheet("요약")
    summary.append(["KB카드 법인영업 데이터 분석 보고서"])
    summary.append([])
    summary.append(["작성일시", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    summary.append(["질문", result.get("question", "조회결과")])
    summary.append(["분석 경로", _get_source_label(result)])
    summary.append(["조회 행 수", len(raw_rows)])
    summary.append(["조회 컬럼 수", len(columns)])
    summary.append([])
    summary.append(["분석 결과", result.get("answer", "") or "분석 결과가 없습니다."])

    data_sheet = wb.create_sheet("상세 데이터")
    data_sheet.freeze_panes = "A2"
    if columns:
        data_sheet.auto_filter.ref = f"A1:{get_column_letter(len(columns))}{len(raw_rows) + 1}"
        data_sheet.append(columns)
        for row in raw_rows:
            data_sheet.append([_excel_cell_value(value) for value in _row_values(row, raw_columns)])
    else:
        data_sheet.append(["조회 데이터가 없습니다."])

    sql_sheet = wb.create_sheet("SQL")
    sql_sheet.append(["실행 SQL"])
    sql_sheet.append([result.get("final_sql", "") or "실행 SQL이 없습니다."])
    wb.save(str(filepath))
    return str(filepath)


def export_to_excel(result: dict) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.table import Table, TableStyleInfo

    _ensure_report_dir()
    question = result.get("question", "조회결과")
    filepath = REPORT_DIR / _make_filename(question, "xlsx")

    _, raw_rows = _raw_result_columns_and_rows(result)
    if len(raw_rows) >= _LARGE_EXCEL_ROW_THRESHOLD:
        return _export_to_excel_write_only(result, filepath)

    columns, rows = _result_columns_and_rows(result)
    answer = result.get("answer", "")
    sql = result.get("final_sql", "")

    wb = Workbook()
    summary = wb.active
    summary.title = "요약"
    summary.sheet_view.showGridLines = False

    title_fill = PatternFill("solid", fgColor="2F5597")
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    section_fill = PatternFill("solid", fgColor="EAF2F8")
    thin_side = Side(style="thin", color="D9E2EC")
    border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    title_font = Font(name="맑은 고딕", size=16, bold=True, color="FFFFFF")
    header_font = Font(name="맑은 고딕", size=10, bold=True, color="1F2937")
    body_font = Font(name="맑은 고딕", size=10, color="1F2937")
    mono_font = Font(name="Consolas", size=9, color="4B5563")

    summary.merge_cells("A1:F1")
    summary["A1"] = "KB카드 법인영업 데이터 분석 보고서"
    summary["A1"].fill = title_fill
    summary["A1"].font = title_font
    summary["A1"].alignment = Alignment(horizontal="center", vertical="center")
    summary.row_dimensions[1].height = 30

    info_rows = [
        ("작성일시", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("질문", question),
        ("분석 경로", _get_source_label(result)),
        ("조회 행 수", len(rows)),
        ("조회 컬럼 수", len(columns)),
    ]
    for row_idx, (label, value) in enumerate(info_rows, start=3):
        summary.cell(row=row_idx, column=1, value=label)
        summary.cell(row=row_idx, column=2, value=value)
        summary.merge_cells(start_row=row_idx, start_column=2, end_row=row_idx, end_column=6)
        for col_idx in range(1, 7):
            cell = summary.cell(row=row_idx, column=col_idx)
            cell.border = border
            cell.font = header_font if col_idx == 1 else body_font
            cell.fill = header_fill if col_idx == 1 else PatternFill(fill_type=None)
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    answer_row = len(info_rows) + 5
    summary.merge_cells(start_row=answer_row, start_column=1, end_row=answer_row, end_column=6)
    summary.cell(row=answer_row, column=1, value="분석 결과")
    summary.cell(row=answer_row, column=1).fill = section_fill
    summary.cell(row=answer_row, column=1).font = header_font
    summary.cell(row=answer_row, column=1).border = border
    summary.cell(row=answer_row, column=1).alignment = Alignment(vertical="center")

    summary.merge_cells(start_row=answer_row + 1, start_column=1, end_row=answer_row + 5, end_column=6)
    answer_cell = summary.cell(row=answer_row + 1, column=1, value=answer or "분석 결과가 없습니다.")
    answer_cell.font = body_font
    answer_cell.border = border
    answer_cell.alignment = Alignment(vertical="top", wrap_text=True)
    for row_idx in range(answer_row + 1, answer_row + 6):
        for col_idx in range(1, 7):
            summary.cell(row=row_idx, column=col_idx).border = border
        summary.row_dimensions[row_idx].height = 24

    for col_idx, width in enumerate([16, 18, 18, 18, 18, 18], start=1):
        summary.column_dimensions[get_column_letter(col_idx)].width = width

    data_sheet = wb.create_sheet("상세 데이터")
    data_sheet.sheet_view.showGridLines = False
    if columns and rows:
        data_sheet.append(columns)
        for cell in data_sheet[1]:
            cell.fill = title_fill
            cell.font = Font(name="맑은 고딕", size=10, bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border

        for row_values in rows:
            data_sheet.append([_excel_cell_value(value) for value in row_values])

        for row in data_sheet.iter_rows(min_row=2, max_row=data_sheet.max_row):
            for cell in row:
                cell.font = body_font
                cell.border = border
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if _is_numeric_value(cell.value):
                    cell.number_format = "#,##0.00" if isinstance(cell.value, float) and not cell.value.is_integer() else "#,##0"
                    cell.alignment = Alignment(horizontal="right", vertical="top", wrap_text=True)
                elif isinstance(cell.value, datetime):
                    cell.number_format = "yyyy-mm-dd hh:mm:ss"
                elif isinstance(cell.value, date):
                    cell.number_format = "yyyy-mm-dd"

        data_sheet.freeze_panes = "A2"
        data_sheet.auto_filter.ref = data_sheet.dimensions
        table_ref = f"A1:{get_column_letter(len(columns))}{len(rows) + 1}"
        table = Table(displayName="ResultTable", ref=table_ref)
        table.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        data_sheet.add_table(table)

        for col_idx, header in enumerate(columns, start=1):
            samples = [header]
            for row_values in rows[:100]:
                samples.append(format_number_for_report(row_values[col_idx - 1]))
            width = min(max(max(len(str(value)) for value in samples) + 2, 10), 42)
            data_sheet.column_dimensions[get_column_letter(col_idx)].width = width
    else:
        data_sheet["A1"] = "조회 데이터가 없습니다."
        data_sheet["A1"].font = body_font
        data_sheet.column_dimensions["A"].width = 32

    sql_sheet = wb.create_sheet("SQL")
    sql_sheet.sheet_view.showGridLines = False
    sql_sheet["A1"] = "실행 SQL"
    sql_sheet["A1"].fill = section_fill
    sql_sheet["A1"].font = header_font
    sql_sheet["A2"] = sql or "실행 SQL이 없습니다."
    sql_sheet["A2"].font = mono_font
    sql_sheet["A2"].alignment = Alignment(vertical="top", wrap_text=True)
    sql_sheet["A2"].border = border
    sql_sheet.column_dimensions["A"].width = 120
    sql_sheet.row_dimensions[2].height = min(360, max(45, (sql.count("\n") + 2) * 15))

    wb.save(str(filepath))
    return str(filepath)


def export_to_text(result: dict) -> str:
    _ensure_report_dir()
    question = result.get("question", "조회결과")
    filepath = REPORT_DIR / _make_filename(question, "txt")
    raw_columns, raw_rows = _raw_result_columns_and_rows(result)
    columns = [str(column) for column in raw_columns]
    answer = result.get("answer", "")
    sql = result.get("final_sql", "")
    with open(filepath, "w", newline="", encoding="utf-8") as f:
        f.write(
            "\n".join(
                [
                    "=" * 70,
                    "  KB카드 법인영업 데이터 분석 보고서",
                    "=" * 70,
                    f"  작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                    f"  질문: {question}",
                    f"  분석 경로: {_get_source_label(result)}",
                    "=" * 70,
                    "",
                ]
            )
        )
        if answer:
            f.write(f"\n[분석 결과]\n{'-' * 70}\n{answer}\n")
        if columns and raw_rows:
            f.write(f"\n[상세 데이터] (총 {len(raw_rows)}건)\n{'-' * 70}\n")
            writer = csv.writer(f, delimiter="\t", lineterminator="\n")
            writer.writerow(columns)
            for row in raw_rows:
                writer.writerow(_row_values(row, raw_columns))
        if sql:
            f.write(f"\n[실행 SQL]\n{'-' * 70}\n{sql}\n")
        f.write(f"\n{'=' * 70}\n  KB카드 법인영업 Text2SQL Agent v10 자동생성 보고서\n{'=' * 70}")
    return str(filepath)


def export_to_csv(result: dict) -> str:
    _ensure_report_dir()
    question = result.get("question", "조회결과")
    filepath = REPORT_DIR / _make_filename(question, "csv")
    raw_columns, raw_rows = _raw_result_columns_and_rows(result)
    columns = [str(column) for column in raw_columns]
    with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        if columns:
            writer.writerow(columns)
        for row in raw_rows:
            writer.writerow(_row_values(row, raw_columns))
    return str(filepath)


def export_all(result: dict) -> dict[str, str]:
    paths = {}
    try:
        paths["word"] = export_to_word(result)
    except ImportError:
        paths["word"] = "(python-docx 패키지 미설치 - pip install python-docx)"
    except Exception as e:
        paths["word"] = f"(Word 생성 실패: {e})"
    try:
        paths["excel"] = export_to_excel(result)
    except ImportError:
        paths["excel"] = "(openpyxl 패키지 미설치 - pip install openpyxl)"
    except Exception as e:
        paths["excel"] = f"(Excel 생성 실패: {e})"
    paths["text"] = export_to_text(result)
    return paths


def _get_source_label(result: dict) -> str:
    tool = result.get("selected_tool", "")
    matched = result.get("matched_query_name", "")
    if tool:
        return f"Tool: {tool}"
    elif matched:
        return f"검증된 쿼리: {matched}"
    elif result.get("question_type") == "direct_sql":
        return "사용자 SQL 실행"
    elif result.get("question_type") == "direct":
        return "직접 답변"
    else:
        return "SQL 자동 생성"
